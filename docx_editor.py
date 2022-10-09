import csv
from docx import Document


def read_csv(csv_file):
    all_data = []
    with open(csv_file, 'r', encoding="utf8") as csvfile:
        reader = csv.reader(csvfile, delimiter=";")
        i = 0
        for row in reader:
            if i == 0:
                headings = row
            else:
                all_data.append(row)
            i += 1
    return headings, all_data


def build_file_name(name_str, replace_dict):
    for placeholder in replace_dict:
        if placeholder in name_str:
            name_str = name_str.replace(placeholder, replace_dict[placeholder])
    return name_str


def create_document(docx_file, headings, data, name_formula, output_dir):
    document = Document(docx_file)
    replace_dict = {x:y for x,y in zip(headings,data)}
    for paragraph in document.paragraphs:
        if paragraph.runs:
            for placeholder in replace_dict:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replace_dict[placeholder])
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        for placeholder in replace_dict:
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, replace_dict[placeholder])
    name = build_file_name(name_formula, replace_dict)
    new_file_name = f"{name}.docx"
    new_file = output_dir + "/" +  new_file_name
    document.save(new_file)
    return new_file_name


def write_documents(docx_template, csv_data, name_formula, output_dir):
    headings, all_data = read_csv(csv_data)
    total = 1
    for data in all_data:
        new_file = create_document(docx_template, headings, data, name_formula, output_dir)
        print(f"{total} - {new_file}")
        total += 1
